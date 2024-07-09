#!/usr/bin/env python
#
# This script can be used for any purpose without limitation subject to the
# conditions at http://www.ccdc.cam.ac.uk/Community/Pages/Licences/v2.aspx
#
# This permission notice and the following statement of attribution must be
# included in all copies or substantial portions of this script.
#
# 2017-08-10: Created by Andy Maloney, the Cambridge Crystallographic Data Centre
# 2020-08-21: made available by the Cambridge Crystallographic Data Centre
#
"""
multi_component_hydrogen_bond_propensity_report.py
 - Performs a multi-component HBP calculation for a given library of co-formers
"""

import sys
import os
import glob
import argparse
import tempfile
import subprocess
import json
import warnings

import matplotlib

matplotlib.use('TkAgg')
import matplotlib.pyplot as plt

from ccdc import io, molecule
from ccdc.diagram import DiagramGenerator
from ccdc.descriptors import CrystalDescriptors

try:
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", category=DeprecationWarning)
        import docxtpl
        from docx.shared import Cm
except ImportError:
    error_message = """
    The python-docx-template templating engine needed by this script could not
    be found. Please run "{} -m pip install docxtpl"
    to try to fix the issue.\nYou may need administrator's rights to do this.
    """.format(sys.executable)
    raise ImportError(error_message)

SCRIPT_DIR = os.path.dirname(__file__)
TEMPLATE_FILENAME = 'multi_component_hydrogen_bond_propensity_report.docx'
TEMPLATE_FILE = os.path.join(SCRIPT_DIR, TEMPLATE_FILENAME)
PAIR_TEMPLATE_FILENAME = 'multi_component_pair_hbp_report.docx'
PAIR_TEMPLATE_FILE = os.path.join(SCRIPT_DIR, PAIR_TEMPLATE_FILENAME)


###############################################################################
def cm2inch(*tupl):
    inch = 2.54
    if isinstance(tupl[0], tuple):
        return tuple(i / inch for i in tupl[0])
    else:
        return tuple(i / inch for i in tupl)


def launch_word_processor(output_file):
    '''This function launches the platform specific word processor
    when not running under continuous integration'''
    if 'TEAMCITY_VERSION' in os.environ:
        return
    if sys.platform == 'win32':
        os.startfile(output_file)
    elif sys.platform.startswith('linux'):
        subprocess.Popen(['xdg-open', output_file])
    else:
        subprocess.Popen(['open', output_file])


def make_diagram(mol, directory):
    # Generates a diagram from a given structure
    molecule_diagram_generator = DiagramGenerator()
    molecule_diagram_generator.settings.line_width = 1.6
    molecule_diagram_generator.settings.font_size = 12
    molecule_diagram_generator.settings.image_height = 300
    img = molecule_diagram_generator.image(mol)
    fname = str(os.path.join(directory, '%s_diagram.png' % mol.identifier))
    if img:
        img.save(fname)
    return fname


def make_mc_chart(dictionary, directory, mol):
    results = [value[0] for key, value in dictionary if isinstance(value[0], float)]
    ymin = min(results)
    ymax = max(results)
    indices = range(1, len(results) + 1)

    fig = plt.figure(figsize=cm2inch(22, 18))
    ax = fig.add_subplot(1, 1, 1)
    color = 'cornflowerblue'
    color1 = 'royalblue'
    ls = ''
    plt.plot(indices, results, marker='D', markersize=10, color=color, ls=ls, markeredgecolor=color1, alpha=0.7)
    plt.axhline(y=0, color='gray')
    plt.xlabel('Co-Former Rank', fontweight='bold', fontsize='12')
    plt.ylabel('Multi-Component Score', fontweight='bold', fontsize='12')
    plt.title('MCHBP screening results', fontweight='bold', fontsize='12')
    ax.axhspan(0.025, ymax + 0.1, facecolor='lightgreen', alpha=0.5)
    ax.axhspan(-0.025, ymin - 0.1, facecolor='lightpink', alpha=0.5)
    ax.axhspan(0.025, -0.025, facecolor='grey', alpha=0.5)
    plt.ylim(ymin - 0.025, ymax + 0.025)
    fname = str(os.path.join(directory, '%s_MC_HBP_plot.png' % mol.identifier))
    plt.savefig(fname, format='png', dpi=600)
    return fname


def add_picture_subdoc(picture_location, docx_template, wd=7):
    # This function adds a picture to the .docx file
    return docxtpl.InlineImage(
        docx_template, image_descriptor=picture_location, width=Cm(wd))


def propensity_calc(crystal, directory):
    # Perform a Hydrogen Bond Propensity calculation

    # Provide settings for the calculation
    settings = CrystalDescriptors.HBondPropensities.Settings()
    settings.working_directory = directory
    settings.hbond_criterion.require_hydrogens = True
    settings.hbond_criterion.path_length_range = (3, 999)

    # Set up the HBP calculator
    hbp = CrystalDescriptors.HBondPropensities(settings)

    # Set up the target structure for the calculation
    hbp.set_target(crystal)

    print(hbp.functional_groups)

    # Generate Training Dataset

    hbp.match_fitting_data(count=500)  # set to >300, preferably 500 for better representation of functional groups

    hbp.analyse_fitting_data()

    for d in hbp.donors:
        print(d.identifier, d.npositive, d.nnegative)
    for a in hbp.acceptors:
        print(a.identifier, a.npositive, a.nnegative)

    # Perform regression
    model = hbp.perform_regression()
    print(model.equation)
    print('Area under ROC curve: {} -- {}'.format(round(model.area_under_roc_curve, 3), model.advice_comment))
    hbp.calculate_propensities()
    propensities = hbp.inter_propensities

    return propensities, hbp.donors, hbp.acceptors


def coordination_scores_calc(crystal, directory):
    # Calculate coordination scores for the target structure

    # Provide settings for the calculation
    settings = CrystalDescriptors.HBondCoordination.Settings()
    settings.working_directory = directory
    settings.hbond_criterion.require_hydrogens = True
    settings.hbond_criterion.path_length_range = (3, 999)

    # Set up the coordination scores calculator
    coordination_calc = CrystalDescriptors.HBondCoordination(settings)
    predictions = coordination_calc.predict(crystal)
    return predictions


def format_scores(scores, das, d_type):
    # Reformat the coordination scores to make report writing easier
    formatted_scores = {}
    for da in das:
        preds = scores.predictions_for_label(da.label, d_type)[1]
        formatted_scores[da.label] = preds
    return formatted_scores


def get_mc_scores(propensities, identifier):
    # Calculates the multi-component scores from the individual HBP calculation
    AA_propensities = []
    BB_propensities = []
    AB_propensities = []
    BA_propensities = []

    for p in propensities:
        t = "%s_d" % p.donor_label.split(" ")[0], "%s_a" % p.acceptor_label.split(" ")[0]
        if '_A_' in t[0] and '_A_' in t[1]:
            AA_propensities.append(p.propensity)
        elif '_B_' in t[0] and '_B_' in t[1]:
            BB_propensities.append(p.propensity)
        elif '_A_' in t[0] and '_B_' in t[1]:
            AB_propensities.append(p.propensity)
        elif '_B_' in t[0] and '_A_' in t[1]:
            BA_propensities.append(p.propensity)
    max_AA = max(AA_propensities) if len(AA_propensities) > 0 else 0.0
    max_BB = max(BB_propensities) if len(BB_propensities) > 0 else 0.0
    max_AB = max(AB_propensities) if len(AB_propensities) > 0 else 0.0
    max_BA = max(BA_propensities) if len(BA_propensities) > 0 else 0.0
    max_list = [max_AA, max_BB, max_AB, max_BA]
    max_keys = ['A:A', 'B:B', 'A:B', 'B:A']
    max_mc = max(max_list[2], max_list[3])
    max_sc = max(max_list[0], max_list[1])

    return [round((max_mc - max_sc), 2),
            max_keys[max_list.index(max(max_list))],
            round(max_mc, 2),
            round(max_list[0], 2),
            round(max_list[1], 2),
            identifier]


def make_pair_file(api_molecule, tempdir, f, i):
    # Creates a file for the api/coformer pair
    with io.MoleculeReader(f) as reader:
        coformer_molecule = reader[0]
        coformer_name = coformer_molecule.identifier
        molecule_pair = make_molecule_pair(api_molecule, coformer_molecule, i)
        molecule_file = os.path.join(tempdir, '%s.mol2' % molecule_pair.identifier)
        with io.MoleculeWriter(molecule_file) as writer:
            writer.write(molecule_pair)
    return molecule_file, coformer_name


def make_molecule_pair(api_molecule, coformer_molecule, i):
    # Creates the multi-component system for each api/coformer pair
    new_file_name = '%s_%d' % (api_molecule.identifier, i)
    molecule_pair = molecule.Molecule(new_file_name)
    molecule_pair.add_molecule(api_molecule)
    molecule_pair.add_molecule(coformer_molecule)
    molecule_pair.normalise_labels()

    atom_list = []
    for atom in molecule_pair.components[0].atoms:
        atom_list.append(atom.label)

    for atom in molecule_pair.atoms:
        if atom.label in atom_list:
            atom.label += '_A'
        else:
            atom.label += '_B'

    return molecule_pair


def pair_output(identifier, propensities, donors, acceptors, coordination_scores, directory):
    # Writes out the output from a single HBP calculation for the multi-component pair

    # This looks for the .docx template that is used to generate the report from
    if os.path.isfile(PAIR_TEMPLATE_FILE):
        docx_template = docxtpl.DocxTemplate(PAIR_TEMPLATE_FILE)
    else:
        print('Error! {} not found!'.format(PAIR_TEMPLATE_FILENAME))
        quit()

    dscores = {}
    ascores = {}

    for d in donors:
        coord_cols = [i for i in range(len(coordination_scores.predictions_for_label(d.label, 'd')[1]))]
        dscores[d.label] = [round((coordination_scores.predictions_for_label(d.label, 'd')[1])[j], 3)
                            for j in coord_cols]

    for a in acceptors:
        ascores[a.label] = [round((coordination_scores.predictions_for_label(a.label, 'a')[1])[k], 3)
                            for k in coord_cols]

    context = {
        'identifier': identifier,
        'propensities': propensities,
        'coord_cols': coord_cols,
        'donors': donors,
        'dscores': dscores,
        'acceptors': acceptors,
        'ascores': ascores
    }
    docx_template.render(context)
    output_file = os.path.join(directory, '%s_pair_output.docx' % identifier)
    docx_template.save(output_file)


def make_mc_report(identifier, results, directory, diagram_file, chart_file):
    # Write the MC-HBP report from the results

    # This looks for the .docx template that is used to generate the report from
    if os.path.isfile(TEMPLATE_FILE):
        docx_template = docxtpl.DocxTemplate(TEMPLATE_FILE)
    else:
        print('Error! {} not found!'.format(TEMPLATE_FILENAME))
        quit()

    # Generate content for the report
    diagram = add_picture_subdoc(diagram_file, docx_template)
    chart = add_picture_subdoc(chart_file, docx_template, wd=18)

    # The context is the information that is given to the template to allow it to be populated
    context = {
        'identifier': str(identifier).split('.')[0],
        'diagram': diagram,
        'chart': chart,
        'results': results
    }

    # Send all the information to the template file then open up the final report
    docx_template.render(context)
    output_file = os.path.join(directory, '%s_MC_HBP_report.docx' % str(identifier).split('.')[0])
    docx_template.save(output_file)

    launch_word_processor(output_file)


def main(structure, work_directory, failure_directory, library, csdrefcode, force_run):
    # This loads up the CSD if a refcode is requested, otherwise loads the structural file supplied
    if csdrefcode:
        try:
            crystal = io.CrystalReader('CSD').crystal(structure)
        except RuntimeError:
            print('Error! %s is not in the database!' % structure)
            quit()
        if io.CrystalReader('CSD').entry(structure).has_disorder and not force_run:
            raise RuntimeError("Disorder can cause undefined behaviour. It is not advisable to run this "
                               "script on disordered entries.\n To force this script to run on disordered entries"
                               " use the flag --force_run_disordered.")
    else:
        crystal = io.CrystalReader(structure)[0]

    # If there are atoms without sites, or there are no atoms in the structure, then HBP cannot be performed
    api_molecule = crystal.molecule
    if not api_molecule.all_atoms_have_sites or len(api_molecule.atoms) == 0:
        print('Error! Not all atoms in %s have sites!' % structure)
        quit()

    # find the coformers and set up the calculations
    coformer_files = glob.glob(os.path.join(library, '*.mol2'))
    tempdir = tempfile.mkdtemp()
    mc_dictionary = {}
    failures = []

    # for each coformer in the library, make a pair file for the api/coformer and run a HBP calculation
    for i, f in enumerate(coformer_files):
        molecule_file, coformer_name = make_pair_file(api_molecule, tempdir, f, i + 1)
        if not io.CrystalReader(f)[0].molecule.is_3d:
            failure_warning = f"Could not run for {coformer_name} no 3d coordinates present."
            failures.append(failure_warning)
            warnings.warn(failure_warning)
            continue
        print(coformer_name)
        crystal_reader = io.CrystalReader(molecule_file)
        crystal = crystal_reader[0]

        directory = os.path.join(os.path.abspath(work_directory), crystal.identifier)
        if os.path.exists(os.path.join(directory, "success.json")):
            with open(os.path.join(directory, "success.json"), "r") as file:
                tloaded = json.load(file)
            mc_dictionary[coformer_name] = tloaded
        else:
            try:
                propensities, donors, acceptors = propensity_calc(crystal, directory)
                coordination_scores = coordination_scores_calc(crystal, directory)
                pair_output(crystal.identifier, propensities, donors, acceptors, coordination_scores, directory)
                with open(os.path.join(directory, "success.json"), "w") as file:
                    tdata = get_mc_scores(propensities, crystal.identifier)
                    json.dump(tdata, file)
                mc_dictionary[coformer_name] = get_mc_scores(propensities, crystal.identifier)
            except RuntimeError as error_message:
                print("Propensity calculation failure for %s!" % coformer_name)
                print(error_message)
                mc_dictionary[coformer_name] = ["N/A", "N/A", "N/A", "N/A", "N/A", crystal.identifier]
                failures.append(f"{coformer_name}: {error_message}")

    # Make sense of the outputs of all the calculations
    mc_hbp_screen = sorted(mc_dictionary.items(), key=lambda e: 0 if e[1][0] == 'N/A' else e[1][0], reverse=True)
    diagram_file = make_diagram(api_molecule, work_directory)
    chart_file = make_mc_chart(mc_hbp_screen, work_directory, api_molecule)
    make_mc_report(structure, mc_hbp_screen, work_directory, diagram_file, chart_file)
    if failure_directory is not None:
        with open(os.path.join(failure_directory, 'failures.txt'), 'w', encoding='utf-8', newline='') as file:
            file.write('\n'.join(map(str, failures)))


if __name__ == '__main__':
    # Set up the necessary arguments to run the script
    # For CSD 2023.1
    ccdc_coformers_dir = os.path.join(
        io.csd_directory(),
        os.pardir, os.pardir,
        'ccdc-software',
        'mercury',
        'molecular_libraries',
        'ccdc_coformers')
    # CSD 2022.3 or earlier
    if not os.path.exists(ccdc_coformers_dir):
        if sys.platform == 'win32':
            ccdc_coformers_dir = os.path.join(
                os.path.dirname(io.csd_directory()),
                'Mercury',
                'molecular_libraries',
                'ccdc_coformers'
            )
        else:
            ccdc_coformers_dir = os.path.join(
                os.path.dirname(io.csd_directory()),
                'molecular_libraries',
                'ccdc_coformers'
            )
    parser = argparse.ArgumentParser(
        formatter_class=argparse.RawDescriptionHelpFormatter,
        description=__doc__)
    parser.add_argument('input_structure', type=str,
                        help='Refcode or mol2 file of the component to be screened')
    parser.add_argument('-d', '--directory', default=os.getcwd(),
                        help='the working directory for the calculation')
    parser.add_argument('-c', '--coformer_library', type=str,
                        help='the directory of the desired coformer library',
                        default=ccdc_coformers_dir)
    parser.add_argument('-f', '--failure_directory', type=str,
                        help='The location where the failures file should be generated')

    parser.add_argument('--force_run_disordered', action="store_true",
                        help='Forces running the script on disordered entries. (NOT RECOMMENDED)', default=False)

    args = parser.parse_args()
    refcode = False
    args.directory = os.path.abspath(args.directory)
    if not os.path.isfile(args.input_structure):
        if len(str(args.input_structure).split('.')) == 1:
            refcode = True
        else:
            parser.error('%s - file not found.' % args.input_structure)
    if not os.path.isdir(args.directory):
        os.makedirs(args.directory)
    if not os.path.isdir(args.coformer_library):
        parser.error('%s - library not found.' % args.coformer_library)

    main(args.input_structure, args.directory, args.failure_directory, args.coformer_library, refcode,
         args.force_run_disordered)
