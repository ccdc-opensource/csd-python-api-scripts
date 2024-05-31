#!/usr/bin/env python
#
# This script can be used for any purpose without limitation subject to the
# conditions at http://www.ccdc.cam.ac.uk/Community/Pages/Licences/v2.aspx
#
# This permission notice and the following statement of attribution must be
# included in all copies or substantial portions of this script.
#
# 2017-08-10: Created by Andy Maloney, the Cambridge Crystallographic Data Centre
# 2020-08-21: made available by the Cambridge Crystallographic Data Centre
# 2022-12-21: Updated by Joanna S. Stevens, the Cambridge Crystallographic Data Centre
#

"""
hydrogen_bond_propensity_report.py
- Writes a .docx report of a hydrogen bond propensity calculation
"""

import argparse
import os
import sys
import subprocess
import csv
import string
import matplotlib

matplotlib.use('Agg')
import matplotlib.pyplot as plt

from ccdc import io
from ccdc.diagram import DiagramGenerator
from ccdc.search import SubstructureSearch, ConnserSubstructure
from ccdc.descriptors import CrystalDescriptors

try:
    import warnings

    with warnings.catch_warnings():
        warnings.simplefilter("ignore", category=DeprecationWarning)
        import docxtpl
        from docx.shared import Cm
except ImportError:
    error_message = """
    The python-docx-template templating engine needed by this script could not
    be found. Please run "{} -m pip install docxtpl"
    to try to fix the issue.\nYou may need administrator's rights to do this.
    """.format(sys.executable)
    raise ImportError(error_message)

SCRIPT_DIR = os.path.dirname(__file__)
TEMPLATE_FILENAME = 'hydrogen_bond_propensity_report.docx'
TEMPLATE_FILE = os.path.join(SCRIPT_DIR, TEMPLATE_FILENAME)


###############################################################################
def make_diagram(mol, directory):
    # Generates a diagram from a given structure
    molecule_diagram_generator = DiagramGenerator()
    molecule_diagram_generator.settings.line_width = 1.6
    molecule_diagram_generator.settings.font_size = 12
    molecule_diagram_generator.settings.image_height = 300
    img = molecule_diagram_generator.image(mol)
    fname = str(os.path.join(directory, '%s_diagram.png' % mol.identifier))
    if img:
        img.save(fname)
    return fname


def generate_component_pics(input_mol, molecule_diagram_generator, directory, docx_template):
    non_ch_molecule_no = 0
    component_diagrams = []
    donor_acceptor_check = []

    for m in input_mol.components:  # generates an image for every molecule in the asymmetric unit
        atoms = m.atoms
        donor_acceptor_check = [atom for atom in atoms if atom.is_donor or atom.is_acceptor]

        if donor_acceptor_check:
            img = molecule_diagram_generator.image(m, highlight_atoms=None, label_atoms=donor_acceptor_check)
            non_ch_molecule_no += 1
            fname = str(os.path.join(directory, '%s_component_%s.png' % (input_mol.identifier, non_ch_molecule_no)))
            if img:
                img.save(fname)
            non_ch_img = add_picture_subdoc(fname, docx_template)
            component_diagrams.append(non_ch_img)
    return component_diagrams


def make_diagram_components(crystal, molecule, directory, docx_template):
    # Generates a diagram for each component in a given structure, with D/A atoms labelled if component has a donor/acceptor
    molecule_diagram_generator = DiagramGenerator()
    molecule_diagram_generator.settings.line_width = 1.6
    molecule_diagram_generator.settings.font_size = 10
    molecule_diagram_generator.settings.image_height = 300
    molecule_diagram_generator.settings.shrink_symbols = False

    asymmetric_molecule = crystal.asymmetric_unit_molecule

    # If no Z and Z' information (e.g. mol2), 'None' is returned for crystal.z_prime
    z_prime_flag = False
    if crystal.z_prime:
        z_prime_flag = True

    # Generates an image for every molecule
    if z_prime_flag and crystal.z_prime < 1:
        component_diagrams = generate_component_pics(molecule, molecule_diagram_generator, directory, docx_template)
    # Generates an image for every molecule in the asymmetric unit
    else:
        component_diagrams = generate_component_pics(asymmetric_molecule, molecule_diagram_generator, directory, docx_template)

    return component_diagrams


def fg_diagram(mol, directory, con):
    # Create highlighted functional group diagrams
    diagram_generator = DiagramGenerator()
    diagram_generator.settings.shrink_symbols = False
    diagram_generator.settings.element_coloring = False
    searcher = SubstructureSearch()
    searcher.add_substructure(ConnserSubstructure(os.path.join(directory, "%s.con" % con)))
    hits = searcher.search(mol)
    selection = hits[0].match_atoms()
    img = diagram_generator.image(mol, highlight_atoms=selection)
    fname = str(os.path.join(directory, '%s.png' % con))
    if img:
        img.save(fname)
    return fname


def add_picture_subdoc(picture_location, docx_template, cm=7):
    # This function adds a picture to the .docx file
    return docxtpl.InlineImage(
        docx_template, image_descriptor=picture_location, width=Cm(cm))


def launch_word_processor(output_file):
    """Open the default application for output_file across platforms."""
    if sys.platform == 'win32':
        os.startfile(output_file)
    elif sys.platform.startswith('linux'):
        subprocess.Popen(['xdg-open', output_file])
    else:
        subprocess.Popen(['open', output_file])


def propensity_calc(crystal, directory, min_donor_coordination, min_acceptor_coordination, fg_count):
    # Perform a Hydrogen Bond Propensity calculation

    # Provide settings for the calculation
    settings = CrystalDescriptors.HBondPropensities.Settings()
    settings.working_directory = directory
    settings.hbond_criterion.require_hydrogens = True
    settings.hbond_criterion.path_length_range = (3, 999)

    # Set up the HBP calculator
    hbp = CrystalDescriptors.HBondPropensities(settings)

    # Set up the target structure for the calculation
    hbp.set_target(crystal)

    print(hbp.functional_groups)

    # Generate Training Dataset
    hbp.match_fitting_data(count=fg_count)  # set to >300, preferably 500 for better representation of functional groups

    hbp.analyse_fitting_data()

    for d in hbp.donors:
        print(d.identifier, d.npositive, d.nnegative)
    for a in hbp.acceptors:
        print(a.identifier, a.npositive, a.nnegative)

    # Perform regression
    model = hbp.perform_regression()

    print(model.equation)
    print('Area under ROC curve: {} -- {}'.format(round(model.area_under_roc_curve, 3), model.advice_comment))

    propensities = hbp.calculate_propensities()

    intra_flag = False
    intra_obs = False
    intra_count = 0
    if len(hbp.intra_propensities) > 0:
        intra_flag = True
        intra_count = len([p for p in propensities if not p.is_intermolecular and p.is_observed])
    if intra_count > 0:
        intra_obs = True

    # Use default coordination cutoffs unless user overrides
    groups = hbp.generate_hbond_groupings(min_donor_prob=min_donor_coordination, min_acceptor_prob=min_acceptor_coordination)

    observed_group = hbp.target_hbond_grouping()

    return hbp.functional_groups, hbp.fitting_data, hbp.donors, hbp.acceptors, model, propensities, \
          intra_flag, groups, observed_group, min_donor_coordination, min_acceptor_coordination, intra_count, intra_obs


def coordination_scores_calc(crystal, directory):
    # Calculate coordination scores for the target structure

    # Provide settings for the calculation
    settings = CrystalDescriptors.HBondCoordination.Settings()
    settings.working_directory = directory

    # Set up the coordination scores calculator
    coordination_calc = CrystalDescriptors.HBondCoordination(settings)
    predictions = coordination_calc.predict(crystal)
    return predictions


def format_scores(scores, das, d_type):
    # Reformat the coordination scores to make report writing easier
    formatted_scores = {}
    for da in das:
        preds = scores.predictions_for_label(da.label, d_type)[1]
        formatted_scores[da.label] = preds
    return formatted_scores


def normalize_molecule(molecule):
    # Normalise bond types for the input structure (important for cifs)
    molecule.assign_bond_types(which='unknown')
    molecule.standardise_aromatic_bonds()
    molecule.standardise_delocalised_bonds()


def chart_output(groups, work_directory, structure):
    # Write out the data points of the HBP chart to a file
    with open(os.path.join(work_directory, '%s_chart_data.csv' % structure), 'w') as outfile:
        csv_writer = csv.writer(outfile)
        csv_writer.writerow(['Mean Propensity', 'Mean Coordination Score', 'Hydrogen Bonds'])
        for group in groups:
            csv_writer.writerow(
                [group.hbond_score,
                 group.coordination_score,
                 '; '.join(['%s - %s' % (g.donor.label, g.acceptor.label) for g in group.hbonds])]
            )

def hbp_landscape(groups, observed_groups, crystal, directory, work_directory, docx_template):
    # Generate the HBP chart

    # 11 colours for alternative networks
    chart_colours = {'#00AEEF': 'lightblue', '#D63868': 'violetred', '#6E3776': 'plum', '#B5B800': 'pear',
                     '#1C5797': 'darkblue', '#FF9A0F': 'orange', '#128A00': 'green', '#D2A106': 'yellow',
                     '#062D56': 'navy', '#00A5A8': 'turquoise', '#EB050C': 'red'}
    hbp_colours = list(chart_colours.keys())

    # If 10 pairs or less in network with highest number of pairs start colour at 0 - static legend
    # otherwise use colours relative to observed (orange, 6th colour, index5) - dynamic legend
    # For half-transparent white edges for triangles, use (1, 1, 1, 0.5) or matplotlib.colors.colorConverter.to_rgba('white', alpha=.5)
    colour_pairs = []
    count_pairs = []
    mpair_flag = False
    obs_pairs = len(observed_groups.hbonds)

    # Check highest number of pairs for alternative networks
    # If no alternative network generated, use number of pairs in observed structure
    try:
        num_alt_pairs = [len(group.hbonds) for group in groups]
        max_alt_pairs = max(num_alt_pairs)
        highest_pairs = max(max_alt_pairs, obs_pairs)
    except:
        highest_pairs = obs_pairs

    # Static legend
    if highest_pairs <= 10 or obs_pairs < 6:
        for i, colour in enumerate(hbp_colours):
            figure_i = plt.scatter([group.hbond_score for group in groups if len(group.hbonds) == i],
                               [abs(group.coordination_score) for group in groups if len(group.hbonds) == i],
                               c=hbp_colours[i], marker='v', s=75, edgecolors=(1, 1, 1, 0.5), linewidth=0.3, clip_on=False)
            colour_pairs.append(i)
            count_pairs.append(len([group for group in groups if len(group.hbonds) == i]))
    # Dynamic legend
    else:
        relative_range = range(5, -6, -1)
        for i, (value, colour) in enumerate(zip(relative_range, hbp_colours)):
            r = obs_pairs - value
            if i == 0:
                figure_i = plt.scatter([group.hbond_score for group in groups if len(group.hbonds) <= r],
                    [abs(group.coordination_score) for group in groups if len(group.hbonds) <= r],
                    c=hbp_colours[i], marker='v', s=75, edgecolors=(1, 1, 1, 0.5), linewidth=0.3, clip_on=False)
                count_pairs.append(len([group for group in groups if len(group.hbonds) <= r]))
                mpair_flag = True
            elif i in range(1, 10):
                figure_i = plt.scatter([group.hbond_score for group in groups if len(group.hbonds) == r],
                                   [abs(group.coordination_score) for group in groups if len(group.hbonds) == r],
                                   c=hbp_colours[i], marker='v', s=75, edgecolors=(1, 1, 1, 0.5), linewidth=0.3, clip_on=False)
                count_pairs.append(len([group for group in groups if len(group.hbonds) == r]))
            elif i == 10:
                figure_i = plt.scatter([group.hbond_score for group in groups if len(group.hbonds) >= r],
                    [abs(group.coordination_score) for group in groups if len(group.hbonds) >= r],
                    c=hbp_colours[i], marker='v', s=75, edgecolors=(1, 1, 1, 0.5), linewidth=0.3, clip_on=False)
                count_pairs.append(len([group for group in groups if len(group.hbonds) >= r]))
                mpair_flag = True
            colour_pairs.append(r)

    # Observed network (plotted after other networks so point on top)
    ax2 = plt.scatter(observed_groups.hbond_score,
                      abs(observed_groups.coordination_score),
                      c='white', marker='o', s=75, edgecolors='black', linewidth=1.6, clip_on=False)

    plt.origin = 'upper'
    plt.xlim(0, 1.0)
    plt.ylim(1.0, 0)
    ax = plt.gca()
    ax.xaxis.tick_top()
    ax.yaxis.tick_left()
    ax.set_xlabel('Mean H-Bond Propensity')
    ax.set_ylabel('Mean H-Bond Co-ordination')
    ax.xaxis.set_label_position('top')
    figure_location = os.path.join(directory, '%s_standard_chart.png' % crystal.identifier)
    plt.savefig(figure_location, bbox_inches='tight', pad_inches=0.0)

    chart = add_picture_subdoc(figure_location, docx_template, cm=12.8)

    diagram_file = make_diagram(crystal.molecule, directory)
    diagram = add_picture_subdoc(diagram_file, docx_template)

    con_files = [f[:-4] for f in os.listdir(work_directory) if f.endswith(".con")]
    fg_pics = [fg_diagram(crystal.molecule, work_directory, con) for con in con_files]
    fg_diagrams = {con: add_picture_subdoc(os.path.join(work_directory, '%s.png' % con), docx_template)
                   for con in con_files}

    return chart, diagram, fg_diagrams, colour_pairs, obs_pairs, count_pairs, mpair_flag


def main(structure, directory, csdrefcode, min_donor_coordination, min_acceptor_coordination, fg_count, noopen=False):
    # This looks for the .docx template that is used to generate the report from
    if os.path.isfile(TEMPLATE_FILE):
        docx_template = docxtpl.DocxTemplate(TEMPLATE_FILE)
    else:
        print('Error! {} not found!'.format(TEMPLATE_FILENAME))
        quit()

    # This loads up the CSD if a refcode is requested, otherwise loads the structural file supplied
    if csdrefcode:
        try:
            crystal = io.CrystalReader('CSD').crystal(str(structure))
        except RuntimeError:
            print('Error! %s is not in the database!' % str(structure))
            quit()
    else:
        crystal = io.CrystalReader(str(structure))[0]

    # If there are atoms without sites, or there are no atoms in the structure, then HBP cannot be performed
    molecule = crystal.molecule
    if not molecule.all_atoms_have_sites or len(molecule.atoms) == 0:
        print('Error! Not all atoms in %s have sites!' % str(structure))
        quit()

    # Bond types need to be standardised
    normalize_molecule(molecule)
    crystal.molecule = molecule

    # Set up a work directory for the HBP files
    work_directory = os.path.join(directory, str(structure).split('.')[0])

    # Get all the necessary data from a HBP calculation
    functional_groups, fitting_data, donors, acceptors, model, propensities, intra_flag, \
    groups, observed_groups, min_donor_coordination, min_acceptor_coordination, intra_count, intra_obs = propensity_calc(crystal, work_directory, min_donor_coordination, min_acceptor_coordination, fg_count)

    # Calculate the coordination scores separately
    coordination_scores = coordination_scores_calc(crystal, work_directory)
    mean_coordination = abs(observed_groups.coordination_score) ## abs() not built into in jinja2 (docx) template

    # Create the HBP chart output as a separate file
    chart_output(groups, work_directory, structure)

    # Set up some dictionaries and fill them with coordination score data, as well as extra information for the report
    dscores = {}
    ascores = {}
    dcoord_bg = {}
    acoord_bg = {}

    for d in donors:
        coord_cols = [i for i in range(len(coordination_scores.predictions_for_label(d.label, 'd')[1]))]
        dscores[d.label] = [round((coordination_scores.predictions_for_label(d.label, 'd')[1])[j], 3)
                            for j in coord_cols]
        dcoord_bg[d.label] = ['FFFFFF' for j in coord_cols]
        coord_value = coordination_scores.predictions_for_label(d.label, 'd')[0]
        if coord_value == dscores[d.label].index(max(dscores[d.label])):
            dcoord_bg[d.label][coord_value] = '7FFF00'
        else:
            dcoord_bg[d.label][coord_value] = 'FF0000'

    for a in acceptors:
        ascores[a.label] = [round((coordination_scores.predictions_for_label(a.label, 'a')[1])[k], 3)
                            for k in coord_cols]
        acoord_bg[a.label] = ['FFFFFF' for j in coord_cols]
        coord_value = coordination_scores.predictions_for_label(a.label, 'a')[0]
        if coord_value == ascores[a.label].index(max(ascores[a.label])):
            acoord_bg[a.label][coord_value] = '7FFF00'
        else:
            acoord_bg[a.label][coord_value] = 'FF0000'

    # Store letters to use as component identifier by using component number as index
    abc = list(string.ascii_uppercase)

    # Generate more information for the report
    coefficients = model.coefficients
    don = list(set(list("%s_d" % p.donor_label.split(" ")[0] for p in propensities)))
    acc = list(set(list("%s_a" % p.acceptor_label.split(" ")[0] for p in propensities)))

    # Generate HBP landscape chart
    chart, diagram, fg_diagrams, colour_pairs, obs_pairs, count_pairs, mpair_flag = hbp_landscape(groups, observed_groups, crystal, directory, work_directory, docx_template)

    # Save 2D diagrams of each chemical component in the asu with non-CH labelled for the SI hbp figure
    component_diagrams = make_diagram_components(crystal, molecule, directory, docx_template)

    # The context is the information that is given to the template to allow it to be populated
    context = {
        # Title page
        'identifier': crystal.identifier,
        'chart': chart,
        'propensities': propensities,
        'intra_flag': intra_flag,
        'coord_cols': coord_cols,
        'donors': donors,
        'acceptors': acceptors,
        'dscores': dscores,
        'dbg': dcoord_bg,
        'ascores': ascores,
        'abg': acoord_bg,
        'diagram': diagram,
        'don': don,
        'acc': acc,
        'fg_diagrams': fg_diagrams,
        'functional_groups': functional_groups,
        'data': fitting_data,
        'len_data': len(fitting_data),
        'coefficients': coefficients,
        'model': model,
        'abc': abc,
        'colour_pairs': colour_pairs,
        'obs_pairs': obs_pairs,
        'count_pairs': count_pairs,
        'min_donor_coordination': min_donor_coordination,
        'min_acceptor_coordination': min_acceptor_coordination,
        'mpair_flag': mpair_flag,
        'component_diagrams': component_diagrams,
        'observed_groups': observed_groups,
        'mean_coordination': mean_coordination,
        'intra_count': intra_count,
        'intra_obs': intra_obs
    }

    # Send all the information to the template file then open up the final report
    docx_template.render(context)
    output_file = os.path.join(directory, '%s_propensity_report.docx' % crystal.identifier)
    docx_template.save(output_file)

    if not noopen:
        launch_word_processor(output_file)
    print('Output file written to %s' % output_file)


if __name__ == '__main__':
    # Set up the necessary arguments to run the script
    parser = argparse.ArgumentParser(
        formatter_class=argparse.RawDescriptionHelpFormatter,
        description=__doc__)
    parser.add_argument('input_structure', type=str,
                        help='Refcode or mol2 file of the component to be screened')
    parser.add_argument('-o', '--output_directory', default=os.getcwd(),
                        help='the working directory for the calculation')
    parser.add_argument('-d', '--min_donor_coordination', default=0.1, type=float,
                        help='Cut-off donor coordination likelihood for HBP chart')
    parser.add_argument('-a', '--min_acceptor_coordination', default=0.1, type=float,
                        help='Cut-off acceptor coordination likelihood for HBP chart')
    parser.add_argument('-c', '--fg_count', default=500, type=int,
                        help='Target functional group count for HBP hits')
    parser.add_argument('-n', '--noopen', action='store_true',
                        help='Do not automatically open the generated output file.')

    args = parser.parse_args()

    refcode = False

    if not os.path.isfile(args.input_structure):
        if len(str(args.input_structure).split('.')) == 1:
            refcode = True
        else:
            parser.error('%s - file not found.' % args.input_structure)
    if not refcode:
        args.directory = os.path.dirname(os.path.abspath(args.input_structure))
    elif not os.path.isdir(args.output_directory):
        os.makedirs(args.output_directory)

    main(args.input_structure, args.output_directory, refcode, args.min_donor_coordination, args.min_acceptor_coordination, args.fg_count, args.noopen)
